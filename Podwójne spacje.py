"""
Przeszukuje wszystkie pliki MS Word (.docx) we wskazanym katalogu i sprawdza je, a tam gdzie odnajdzie
podwójną spację (podwójny biały znak jest niepożądany przy redakcji tekstu), to zamienia ją na pojedynczą spację.
Na koniec swojego działania, zapisuje zmiany.

Każdy zmodyfikowany plik, aby odróżniał się od tych niezmodyfikowanych (aby było wiadomo, że akurat w tym pliku
coś było robione, ma zmienianą nazwę z przedrostkiem [mod] oraz sygnaturę w postaci daty i godziny wraz z sekundami.

"""


import os
from docx import Document
import platform
import time


def remove_double_spaces_in_docx(file_path):
    doc = Document(file_path)
    total_double_spaces = 0  # Zmienna do przechowywania łącznej liczby podwójnych spacji
    for paragraph in doc.paragraphs:
        if '  ' in paragraph.text:
            total_double_spaces += paragraph.text.count('  ')
            paragraph.text = ' '.join(paragraph.text.split())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '  ' in cell.text:
                    total_double_spaces += cell.text.count('  ')
                    cell.text = ' '.join(cell.text.split())

    # Tylko jeśli znaleziono podwójne spacje, zapisz plik z nową nazwą
    if total_double_spaces > 0:
        current_time = time.strftime("%d-%m-%Y_%H-%M-%S")
        new_file_name = "[mod] " + os.path.splitext(os.path.basename(file_path))[0] + "_" + current_time + ".docx"
        new_file_path = os.path.join(os.path.dirname(file_path), new_file_name)
        doc.save(new_file_path)
        os.remove(file_path)  # Usuń oryginalny plik
        return total_double_spaces, new_file_name
    return total_double_spaces, None


def process_directory(directory):
    start_time = time.time()
    modified_files = 0
    total_double_spaces_replaced = 0

    for filename in os.listdir(directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(directory, filename)
            double_spaces_in_file, new_file_name = remove_double_spaces_in_docx(file_path)
            if double_spaces_in_file > 0:
                modified_files += 1
                total_double_spaces_replaced += double_spaces_in_file
                print(f"\nPrzetworzono plik: {new_file_name}, liczba podwójnych spacji: {double_spaces_in_file}")

    end_time = time.time()
    total_time = end_time - start_time

    print(f"\n\nPodsumowanie:")
    print(f"Liczba zmodyfikowanych plików: {modified_files}")
    print(f"Całkowity czas operacji: {total_time:.2f} sekund")
    print(f"Łącznie zamieniono podwójnych spacji na pojedyncze: {total_double_spaces_replaced}")


def menu():
    while True:
        print("\n\nCo Kasiu robimy dalej?")
        print("1. Szukamy tutaj")
        print("2. Szukamy w nowym katalogu")
        try:
            choice = int(input())
            if 1 <= choice <= 2:
                return choice
        except ValueError:
            pass


def main_loop():
    directory = input("Podaj ścieżkę do folderu z plikami .docx: ")
    while True:
        process_directory(directory)

        choice = menu()
        if choice == 1:
            continue
        elif choice == 2:
            directory = input("Podaj ścieżkę do nowego katalogu: ")


if __name__ == "__main__":
    main_loop()

    # if platform.system() != 'Windows':
    #     input("Press any button to continue...")  # Dla Unix/Linux i MacOS

    # Oczekiwanie na naciśnięcie dowolnego klawisza przed zamknięciem
    if platform.system() == 'Windows':
        os.system('pause')  # Tylko dla Windows
    else:
        input("Press any button to continue...")  # Dla Unix/Linux i MacOS
