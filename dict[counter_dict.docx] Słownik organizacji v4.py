import docx
import os
from collections import Counter
from nltk.tokenize import word_tokenize
import re
import datetime

def find_docx_files(folder_path):
    return [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith('.docx')]

def read_docx(file_path):
    doc = docx.Document(file_path)
    return ' '.join([para.text for para in doc.paragraphs])

def process_text(text):
    text = re.sub(r'[^\w\s]', ' ', text)  # Zamiana znaków interpunkcyjnych na spacje
    tokens = word_tokenize(text)
    filtered_tokens = [word for word in tokens if len(word) > 1]  # Filtruje jednoliterowe tokeny
    return filtered_tokens

def get_keywords_and_frequent_words(tokens):
    count = Counter(tokens)
    frequent_words = [word for word, freq in count.most_common(180)]  # 20 najczęstszych słów
    keywords = list(set([word for word in tokens if word.istitle()]))[:80]  # 20 słów zaczynających się z dużej litery
    return keywords, frequent_words

def save_results(folder_path, original_file_name, keywords, frequent_words):
    timestamp = datetime.datetime.now().strftime("%d-%m-%Y_%H-%M-%S")
    file_name = f"[counter_dict] {original_file_name}_{timestamp}.docx"
    doc = docx.Document()
    doc.add_heading('Słowa kluczowe i nazwy', level=1)
    doc.add_paragraph(', '.join(keywords))
    doc.add_heading('Najczęściej powtarzające się słowa', level=1)
    doc.add_paragraph(', '.join(frequent_words))
    doc.save(os.path.join(folder_path, file_name))
    print(f"Zapisano wyniki w {os.path.join(folder_path, file_name)}")

folder_path = input("Podaj ścieżkę do folderu: ")
file_paths = find_docx_files(folder_path)

for file_path in file_paths:
    text = read_docx(file_path)
    tokens = process_text(text)
    keywords, frequent_words = get_keywords_and_frequent_words(tokens)
    original_file_name = os.path.splitext(os.path.basename(file_path))[0]
    save_results(folder_path, original_file_name, keywords, frequent_words)
